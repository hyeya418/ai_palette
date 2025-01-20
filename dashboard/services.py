import os
import logging
from django.conf import settings
from langchain.docstore.document import Document as LangchainDocument
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

# Django settings에서 API 키를 가져옵니다
os.environ["OPENAI_API_KEY"] = settings.OPENAI_API_KEY

def process_uploaded_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    documents = []

    if ext == '.pdf':
        loader = PyPDFLoader(file_path)
        documents = loader.load()
    elif ext == '.docx':
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = [LangchainDocument(page_content=text, metadata={"source": os.path.basename(file_path)})]
    elif ext in ['.xls', '.xlsx']:
        workbook = load_workbook(file_path)
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                text = "\t".join([str(cell.value) for cell in row])
                documents.append(LangchainDocument(page_content=text, metadata={"source": f"{os.path.basename(file_path)} - {sheet.title}"}))
    elif ext == '.hwp':
        # 한글 파일 로딩 로직 추가 (필요한 라이브러리 설치 필요)
        pass
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            documents = [LangchainDocument(page_content=text, metadata={"source": os.path.basename(file_path)})]
    else:
        raise ValueError("Unsupported file type: {}".format(ext))

    return documents

def save_faiss_index(documents, doc_no):
    embeddings = OpenAIEmbeddings()
    vectors = FAISS.from_documents(documents, embeddings)
    index_file_path = f'vectors/faiss_index_{doc_no}.index'  # 보고서 번호를 인덱스 파일 경로에 반영

    logger.info(f"Saving FAISS index to: {index_file_path}")
    vectors.save_local(index_file_path)
    return index_file_path


def load_faiss_index(index_file_path):
    embeddings = OpenAIEmbeddings()
    try:
        vectors = FAISS.load_local(index_file_path, embeddings, allow_dangerous_deserialization=True)
        logger.info(f"Successfully loaded FAISS index from: {index_file_path}")
        return vectors
    except Exception as e:
        logger.error(f"Error loading FAISS index from {index_file_path}: {str(e)}")
        raise
